"""Minimal Markdown → .docx converter for the Lumen progress report.

Handles: ATX headings, fenced code blocks, GitHub pipe tables, bullet lists,
horizontal rules, **bold**, `inline code`, and [text](link) (rendered as text).
Not a general Markdown engine — just enough for our report files.
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")


def _strip_inline(text: str) -> str:
    """Resolve links to their label; leave bold/code markers for the runner."""
    return LINK_RE.sub(r"\1", text)


def add_runs(paragraph, text: str):
    """Add text to a paragraph, honoring **bold** and `code` spans."""
    text = _strip_inline(text)
    # Split on bold and code while keeping the delimiters' content.
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in token_re.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_table(doc: Document, rows: list[list[str]]):
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell in enumerate(header):
        para = table.rows[0].cells[i].paragraphs[0]
        add_runs(para, cell)
        for r in para.runs:
            r.bold = True
    for row in body:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            if i < len(cells):
                add_runs(cells[i].paragraphs[0], cell)


def is_table_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[:\- |]+\|?\s*", line)) and "-" in line


def parse_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def convert(md_path: str, docx_path: str):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block
        if line.lstrip().startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # Table (header row followed by a separator row)
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            rows = [parse_row(line)]
            i += 2
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(parse_row(lines[i]))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Horizontal rule
        if re.fullmatch(r"\s*---+\s*", line):
            doc.add_paragraph().add_run("─" * 40).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        m = re.match(r"(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            heading.clear()
            add_runs(heading, m.group(2))
            i += 1
            continue

        # Bullets
        m = re.match(r"(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, m.group(2))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
